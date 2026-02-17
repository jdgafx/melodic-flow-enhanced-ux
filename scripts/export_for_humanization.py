#!/usr/bin/env python3
"""
AMP Marketing — Export All Website Copy for Humanization
=========================================================
Extracts ALL visible text from every page on the AMP Marketing website,
organized with clear section markers so humanized content can be mapped
back to the exact correct locations.

Output:
  - scripts/amp_marketing_full_copy.docx  (Word format for AISEO upload)
  - scripts/amp_marketing_full_copy.txt   (Plain text backup)
  - scripts/content_map.json             (Machine-readable map for re-import)

Usage: python scripts/export_for_humanization.py
"""

import json
import os
import re
import sys

# ── Page Definitions ──────────────────────────────────────────────────────────
# Each entry: (page_name, file_path, display_name)
PAGES = [
    ("homepage", "src/app/page.tsx", "Homepage"),
    ("about", "src/app/about/page.tsx", "About"),
    ("contact", "src/app/contact/page.tsx", "Contact"),
    ("pricing", "src/app/pricing/page.tsx", "Pricing"),
    ("services", "src/app/services/page.tsx", "Services Overview"),
    ("ai-chatbot", "src/app/services/ai-chatbot/page.tsx", "AI Chatbot"),
    ("ai-voice", "src/app/services/ai-voice/page.tsx", "AI Voice"),
    ("ad-copy", "src/app/services/ad-copy/page.tsx", "Ad Copy"),
    (
        "email-automation",
        "src/app/services/email-automation/page.tsx",
        "Email Automation",
    ),
    ("google-business", "src/app/services/google-business/page.tsx", "Google Business"),
    ("landing-pages", "src/app/services/landing-pages/page.tsx", "Landing Pages"),
    ("lead-funnel", "src/app/services/lead-funnel/page.tsx", "Lead Funnel"),
    ("review-response", "src/app/services/review-response/page.tsx", "Review Response"),
    ("seo-content", "src/app/services/seo-content/page.tsx", "SEO Content"),
    ("social-media", "src/app/services/social-media/page.tsx", "Social Media"),
]


def decode_html_entities(val):
    val = val.replace("&apos;", "'").replace("&quot;", '"').replace("&amp;", "&")
    val = val.replace("\\&apos;", "'").replace("\\&quot;", '"')
    return val


def extract_double_quoted_string(content, start_pos):
    """Extract a full double-quoted string handling escapes properly."""
    if start_pos >= len(content) or content[start_pos] != '"':
        return None, start_pos
    i = start_pos + 1
    result = []
    while i < len(content):
        ch = content[i]
        if ch == "\\" and i + 1 < len(content):
            result.append(content[i + 1])
            i += 2
        elif ch == '"':
            return "".join(result), i + 1
        else:
            result.append(ch)
            i += 1
    return None, start_pos


def extract_text_blocks(content):
    blocks = []

    dq_field_pattern = re.compile(
        r'(?:desc|description|quote|story|solution|result|testimonial|answer|question):\s*"'
    )
    for m in dq_field_pattern.finditer(content):
        quote_start = m.end() - 1
        val, _ = extract_double_quoted_string(content, quote_start)
        if val and len(val) > 15:
            val = decode_html_entities(val)
            blocks.append(("data", val, m.start()))

    bullet_pattern = re.compile(r'^\s*"', re.MULTILINE)
    for m in bullet_pattern.finditer(content):
        quote_start = m.end() - 1
        val, end_pos = extract_double_quoted_string(content, quote_start)
        if not val or len(val) < 20:
            continue
        rest = content[end_pos : end_pos + 5].strip()
        if not rest or rest[0] in ",]":
            if (
                not val.startswith("/")
                and not val.endswith(".svg")
                and not val.startswith("http")
            ):
                val = decode_html_entities(val)
                blocks.append(("bullet", val, m.start()))

    jsx_text_pattern = re.compile(r">([^<{]+)<")
    for m in jsx_text_pattern.finditer(content):
        text = m.group(1).strip()
        text = decode_html_entities(text)
        if len(text) > 25 and not text.startswith("{") and not text.startswith("$"):
            blocks.append(("jsx", text, m.start()))

    meta_block = re.search(r"export const metadata\s*=\s*\{([\s\S]+?)\};", content)
    if meta_block:
        meta_text = meta_block.group(1)
        for field in ["title", "description"]:
            fm = re.search(rf'{field}:\s*"', meta_text)
            if fm:
                val, _ = extract_double_quoted_string(meta_text, fm.end() - 1)
                if val and len(val) > 15 and "AMP Marketing" not in val:
                    blocks.append(("meta", val, meta_block.start() + fm.start()))

    faq_q_pattern = re.compile(r'q:\s*"')
    for m in faq_q_pattern.finditer(content):
        val, _ = extract_double_quoted_string(content, m.end() - 1)
        if val:
            blocks.append(("faq_q", decode_html_entities(val), m.start()))

    faq_a_pattern = re.compile(r'a:\s*"')
    for m in faq_a_pattern.finditer(content):
        val, _ = extract_double_quoted_string(content, m.end() - 1)
        if val and len(val) > 15:
            blocks.append(("faq_a", decode_html_entities(val), m.start()))

    seen = set()
    unique_blocks = []
    for btype, text, pos in sorted(blocks, key=lambda x: x[2]):
        normalized = text.strip()[:80]
        if normalized not in seen:
            seen.add(normalized)
            unique_blocks.append((btype, text))

    return unique_blocks


def format_section_label(btype):
    """Human-readable label for block type."""
    labels = {
        "meta": "SEO META",
        "jsx": "PAGE COPY",
        "data": "CONTENT BLOCK",
        "bullet": "BULLET POINT",
        "faq_q": "FAQ QUESTION",
        "faq_a": "FAQ ANSWER",
    }
    return labels.get(btype, "TEXT")


def build_export():
    """Build the complete export file."""
    os.chdir(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

    content_map = {}  # For machine-readable re-import
    txt_lines = []  # For .txt output

    txt_lines.append("=" * 80)
    txt_lines.append("AMP MARKETING — COMPLETE WEBSITE COPY")
    txt_lines.append("Export for Humanization")
    txt_lines.append("=" * 80)
    txt_lines.append("")
    txt_lines.append("INSTRUCTIONS:")
    txt_lines.append("- Each section below is marked with ===PAGE=== headers")
    txt_lines.append("- Within each page, text blocks are numbered [BLOCK-XX]")
    txt_lines.append(
        "- When humanizing, KEEP the ===PAGE=== and [BLOCK-XX] markers intact"
    )
    txt_lines.append("- Only change the text BETWEEN the markers")
    txt_lines.append("- Save the humanized version and we'll map it back automatically")
    txt_lines.append("")
    txt_lines.append("=" * 80)
    txt_lines.append("")

    total_blocks = 0
    total_words = 0

    for page_id, file_path, display_name in PAGES:
        if not os.path.exists(file_path):
            print(f"  ⚠️  Skipping {file_path} (not found)")
            continue

        with open(file_path, "r", encoding="utf-8") as f:
            raw_content = f.read()

        blocks = extract_text_blocks(raw_content)

        if not blocks:
            print(f"  ⚠️  No text blocks found in {file_path}")
            continue

        # Page header
        txt_lines.append(f"===PAGE: {display_name} ({file_path})===")
        txt_lines.append("")

        page_blocks = []
        block_num = 1

        for btype, text in blocks:
            block_id = f"{page_id}-{block_num:02d}"
            label = format_section_label(btype)

            txt_lines.append(f"[BLOCK-{block_id}] ({label})")
            txt_lines.append(text)
            txt_lines.append(f"[/BLOCK-{block_id}]")
            txt_lines.append("")

            page_blocks.append(
                {
                    "id": block_id,
                    "type": btype,
                    "original": text,
                }
            )

            word_count = len(text.split())
            total_words += word_count
            total_blocks += 1
            block_num += 1

        content_map[file_path] = {
            "page_id": page_id,
            "display_name": display_name,
            "blocks": page_blocks,
        }

        txt_lines.append(f"===END PAGE: {display_name}===")
        txt_lines.append("")
        txt_lines.append("")

        print(
            f"  ✅ {display_name}: {len(blocks)} blocks, {sum(len(b[1].split()) for b in blocks)} words"
        )

    # Summary footer
    txt_lines.append("=" * 80)
    txt_lines.append(
        f"TOTAL: {total_blocks} text blocks across {len(content_map)} pages"
    )
    txt_lines.append(f"TOTAL WORDS: ~{total_words:,}")
    txt_lines.append("=" * 80)

    # ── Write .txt ────────────────────────────────────────────────────────────
    txt_path = "scripts/amp_marketing_full_copy.txt"
    os.makedirs("scripts", exist_ok=True)
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("\n".join(txt_lines))
    print(f"\n📄 Text file: {txt_path}")

    # ── Write .json (for machine re-import) ───────────────────────────────────
    json_path = "scripts/content_map.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(content_map, f, indent=2, ensure_ascii=False)
    print(f"📋 Content map: {json_path}")

    # ── Write .docx ───────────────────────────────────────────────────────────
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading("AMP Marketing — Complete Website Copy", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("Export for Humanization via AISEO.ai")
        doc.add_paragraph("")

        # Instructions
        inst = doc.add_paragraph()
        inst.add_run("INSTRUCTIONS:\n").bold = True
        inst.add_run("• Each section is marked with PAGE headers\n")
        inst.add_run("• Text blocks are numbered [BLOCK-XX]\n")
        inst.add_run("• When humanizing, KEEP the PAGE and BLOCK markers intact\n")
        inst.add_run("• Only change the text BETWEEN the markers\n")
        inst.add_run(
            "• Save the humanized version and we'll map it back automatically\n"
        )

        doc.add_page_break()

        for file_path, page_data in content_map.items():
            # Page header
            h = doc.add_heading(f"PAGE: {page_data['display_name']}", level=1)
            p = doc.add_paragraph()
            run = p.add_run(f"Source: {file_path}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)

            for block in page_data["blocks"]:
                label = format_section_label(block["type"])

                # Block marker
                marker_p = doc.add_paragraph()
                marker_run = marker_p.add_run(f"[BLOCK-{block['id']}] ({label})")
                marker_run.bold = True
                marker_run.font.size = Pt(10)
                marker_run.font.color.rgb = RGBColor(79, 70, 229)  # indigo

                # Content
                content_p = doc.add_paragraph(block["original"])
                content_p.style.font.size = Pt(11)

                # End marker
                end_p = doc.add_paragraph()
                end_run = end_p.add_run(f"[/BLOCK-{block['id']}]")
                end_run.font.size = Pt(10)
                end_run.font.color.rgb = RGBColor(79, 70, 229)

            doc.add_page_break()

        # Summary
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(
            f"Total: {total_blocks} text blocks across {len(content_map)} pages"
        )
        doc.add_paragraph(f"Total words: ~{total_words:,}")

        docx_path = "scripts/amp_marketing_full_copy.docx"
        doc.save(docx_path)
        print(f"📝 Word file: {docx_path}")

    except ImportError:
        print("⚠️  python-docx not installed, skipping .docx generation")
        print("   Install with: pip install python-docx")

    print(
        f"\n🎯 DONE! {total_blocks} blocks, ~{total_words:,} words extracted from {len(content_map)} pages"
    )
    print(f"\n📌 Next steps:")
    print(
        f"   1. Upload scripts/amp_marketing_full_copy.docx to AISEO 'Make it Undetectable'"
    )
    print(f"   2. Download the humanized version")
    print(f"   3. Run: python scripts/apply_humanization.py <humanized_file>")


if __name__ == "__main__":
    build_export()
